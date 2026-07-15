"""
CV Processor — Extract text from PDF and DOCX files.
Dilengkapi perbaikan encoding/mojibake untuk mendukung berbagai bahasa
(Indonesia, Arab, Mandarin, Jepang, dll).
"""

import io
import unicodedata
from pathlib import Path
import pdfplumber
import ftfy
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def _clean_text(text: str) -> str:
    if not text:
        return text
    fixed = ftfy.fix_text(text)
    normalized = unicodedata.normalize("NFC", fixed)
    return normalized


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    raw_text = "\n\n".join(text_parts)
    cleaned = _clean_text(raw_text)

    if len(cleaned.strip()) < 50:
        try:
            ocr_text = _extract_text_via_gemini_vision(file_bytes)
            if ocr_text and len(ocr_text.strip()) >= 50:
                return _clean_text(ocr_text)
        except Exception:
            pass  # fallback gagal - kembalikan hasil asli (mungkin kosong)

    return cleaned


def _extract_text_via_gemini_vision(file_bytes: bytes) -> str:
    """
    Convert halaman PDF jadi gambar, lalu pakai Gemini vision untuk
    ekstrak teksnya. Dipakai sebagai fallback untuk PDF hasil scan/foto
    atau export dari tool desain (Canva, dll) yang tidak punya layer
    teks digital.
    """
    import config
    if not config.is_gemini_configured():
        return ""

    from pdf2image import convert_from_bytes
    from google.genai import types

    images = convert_from_bytes(file_bytes, dpi=200)
    client = config.get_gemini_client()

    all_text = []
    for img in images[:5]:  # batasi 5 halaman - CV umumnya 1-2 halaman
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()

        response = client.models.generate_content(
            model=config.GEMINI_MODEL,
            contents=[
                types.Part.from_bytes(data=img_bytes, mime_type="image/png"),
                "Ekstrak SEMUA teks yang terlihat di gambar ini apa adanya, "
                "pertahankan urutan dan struktur (baris baru, poin-poin). "
                "Ini adalah halaman CV/resume. JANGAN tambahkan komentar, "
                "penjelasan, atau kalimat pembuka/penutup - HANYA teks hasil ekstraksi.",
            ],
        )
        if response.text:
            all_text.append(response.text)

    return "\n\n".join(all_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
    except PackageNotFoundError:
        raise ValueError(
            "File ini terdeteksi sebagai format .doc lama (Word 97-2003), "
            "bukan .docx. Silakan convert ke .docx atau PDF terlebih dahulu, "
            "misalnya lewat Google Docs (buka file lalu 'Save as .docx')."
        )

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    raw_text = "\n".join(text_parts)
    return _clean_text(raw_text)


def extract_cv_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")

    if not text or not text.strip():
        raise ValueError(
            "Tidak ada teks yang berhasil diekstrak dari file ini, walau sudah "
            "dicoba fallback pembacaan gambar (OCR). Kemungkinan file rusak, "
            "kualitas scan terlalu buruk, atau benar-benar kosong."
        )
    return text


def get_file_info(file_bytes: bytes, filename: str) -> dict:
    ext = Path(filename).suffix.lower()
    size_mb = len(file_bytes) / (1024 * 1024)
    info = {
        "filename": filename,
        "format": ext.upper().replace(".", ""),
        "size_mb": round(size_mb, 2),
        "size_bytes": len(file_bytes),
    }
    if ext == ".pdf":
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                info["pages"] = len(pdf.pages)
        except Exception:
            info["pages"] = "Unknown"
    elif ext in (".docx", ".doc"):
        try:
            doc = Document(io.BytesIO(file_bytes))
            info["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
        except PackageNotFoundError:
            info["paragraphs"] = "N/A (format .doc lama)"
        except Exception:
            info["paragraphs"] = "Unknown"
    return info


def validate_cv_file(file_bytes: bytes, filename: str, max_size_mb: int = 100) -> tuple[bool, str]:
    ext = Path(filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        return False, f"Format {ext} tidak didukung. Gunakan PDF atau DOCX."

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return False, f"Ukuran file {size_mb:.1f}MB melebihi batas {max_size_mb}MB."

    if len(file_bytes) == 0:
        return False, "File kosong atau gagal terbaca saat upload."

    return True, ""
