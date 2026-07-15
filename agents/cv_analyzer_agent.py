"""
CV Analyzer Agent — Reviews CV content and generates feedback.
Provides ATS score, improvement suggestions, and generates ATS-friendly CV.
"""

import io
from typing import Optional
import config  # noqa: keep single import
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


REVIEW_PROMPT = """Kamu adalah CV Review Expert dengan pengalaman 10+ tahun di bidang HR dan recruitment. 
Tugasmu adalah menganalisis CV user dan memberikan feedback komprehensif.

Berikan output dalam format berikut (gunakan heading yang sama persis):

## 📊 ATS Score: [score]/100

## ✅ Kelebihan CV
- [point 1]
- [point 2]
...

## ⚠️ Area yang Perlu Diperbaiki
- [point 1]
- [point 2]
...

## 💡 Saran Perbaikan
- [saran spesifik 1]
- [saran spesifik 2]
...

## 🔑 Keywords yang Terdeteksi
[list keywords/skills yang ditemukan di CV]

## 📝 Ringkasan Profil
[ringkasan singkat profil kandidat berdasarkan CV]

Jawab dalam Bahasa Indonesia. Berikan feedback yang spesifik dan actionable."""


ATS_CV_PROMPT = """Kamu adalah CV Writer Expert profesional dan spesialis optimasi ATS (Applicant Tracking System).
Tugasmu adalah mengubah CV asli user menjadi CV profesional yang sangat ATS-friendly, rapi, terstruktur, dan memiliki daya jual tinggi bagi HRD.

Ikuti petunjuk struktur dan pemformatan berikut dengan sangat ketat:

### 1. FORMAT & STRUKTUR CV
Gunakan susunan bagian/section berikut secara berurutan dengan heading bertanda `## `:

## [NAMA LENGKAP KANDIDAT]
[Email] | [Nomor Telepon] | [Kota/Kabupaten, Provinsi] | [Link LinkedIn]

## RINGKASAN PROFIL
Tulis ringkasan profesional singkat (3-4 kalimat) yang merangkum keahlian utama, pengalaman relevan, nilai jual unik kandidat, dan apa yang ingin mereka capai secara profesional. Buat agar sangat menarik bagi perekrut.

## PENGALAMAN KERJA
Format setiap riwayat pekerjaan secara konsisten sebagai berikut:
**[Nama Jabatan/Posisi]** | [Nama Perusahaan] | [Bulan Tahun Mulai] – [Bulan Tahun Selesai/Sekarang]
- Gunakan poin-poin bullet (`- `).
- Setiap poin harus dimulai dengan Kata Kerja Aksi yang kuat (Action Verbs), contoh: *Merancang, Mengimplementasikan, Memimpin, Meningkatkan, Mengoptimalkan, Menganalisis*.
- Gunakan formula XYZ untuk menulis pencapaian: "Berhasil mencapai [X], diukur dengan [Y], dengan melakukan [Z]". Usahakan memasukkan angka/metrik konkret (misal: persentase kenaikan, efisiensi waktu, jumlah user) untuk memberikan dampak visual yang kuat.

## PENDIDIKAN
Format setiap riwayat pendidikan secara konsisten sebagai berikut:
**[Nama Gelar/Jurusan]** | [Nama Universitas/Sekolah] | [Tahun Kelulusan]
- Tambahkan info relevan jika ada (misal: IPK jika di atas 3.0/4.0, atau pencapaian akademis penting).

## KEAHLIAN (SKILLS)
Kelompokkan keahlian ke dalam kategori agar mudah dibaca oleh ATS dan HRD, contoh:
- **Keahlian Teknis (Hard Skills)**: [Daftar keahlian teknis utama]
- **Alat & Teknologi (Tools)**: [Software/Tools/Bahasa Pemrograman yang dikuasai]
- **Keahlian Interpersonal (Soft Skills)**: [Daftar soft skills yang relevan]

## SERTIFIKASI & PROJEK (Opsional)
Jika ada sertifikasi atau projek penting di CV asli, format sebagai berikut:
**[Nama Sertifikasi/Projek]** | [Penerbit/Penyelenggara/Tahun]

### 2. ATURAN PENULISAN (RULES)
1. JANGAN gunakan emoji dekoratif di dalam konten CV karena dapat membingungkan sistem parse ATS.
2. Gunakan tata bahasa profesional (Bahasa Indonesia baku atau Bahasa Inggris profesional, sesuaikan dengan bahasa utama pada CV asli).
3. HANYA keluarkan teks CV hasil optimasi saja dari baris pertama hingga terakhir. JANGAN berikan kalimat pembuka ("Berikut adalah...", "Tentu, ini CV...", dll.) atau kalimat penutup."""


def review_cv(cv_text: str) -> dict:
    """
    Analyze CV and return structured feedback.

    Returns dict with:
    - "feedback": AI-generated feedback markdown
    - "available": whether OpenAI is configured
    """
    if not config.is_gemini_configured():
        return {
            "feedback": None,
            "available": False,
        }

    try:
        client = config.get_gemini_client()
        prompt = f"{REVIEW_PROMPT}\n\nBerikut adalah CV yang perlu di-review:\n\n{cv_text}"
        response = client.models.generate_content(
            model=config.GEMINI_MODEL,
            contents=prompt,
        )
        return {
            "feedback": response.text,
            "available": True,
        }
    except Exception as e:
        return {
            "feedback": f"Error: {str(e)}",
            "available": True,
        }


def generate_ats_cv(cv_text: str, job_info: Optional[dict] = None, language: str = "auto") -> dict:
    """
    Generate an ATS-friendly version of the CV, optionally tailored to a target job.

    Returns dict with:
    - "ats_text": improved CV content as plain text
    - "available": whether Gemini is configured
    """
    if not config.is_gemini_configured():
        return {"ats_text": None, "available": False}

    try:
        client = config.get_gemini_client()
        prompt = f"{ATS_CV_PROMPT}\n\nCV Asli:\n\n{cv_text}"

        if language == "id":
            prompt += (
                "\n\n[INSTRUKSI BAHASA]\n"
                "HASILKAN CV DALAM BAHASA INDONESIA, apa pun bahasa asli CV di atas."
            )
        elif language == "en":
            prompt += (
                "\n\n[LANGUAGE INSTRUCTION]\n"
                "OUTPUT THE CV IN ENGLISH, regardless of the original CV's language."
            )
        
        if job_info:
            target_pos = job_info.get("job_title", "N/A")
            target_company = job_info.get("company_name", "N/A")
            target_desc = job_info.get("job_description", "N/A")
            prompt += (
                f"\n\n[INFO LOWONGAN SASARAN]\n"
                f"Posisi: {target_pos}\n"
                f"Perusahaan: {target_company}\n"
                f"Deskripsi Lowongan:\n{target_desc[:1500]}\n\n"
                f"TUGAS PENTING: Optimalkan CV di atas agar sangat relevan dengan lowongan sasaran ini. "
                f"Sesuaikan pilihan kata kunci (keywords), highlight pengalaman kerja yang paling cocok, dan sesuaikan ringkasan profil agar mencerminkan kecocokan kuat dengan posisi ini."
            )
            
        response = client.models.generate_content(
            model=config.GEMINI_MODEL,
            contents=prompt,
        )
        return {
            "ats_text": response.text,
            "available": True,
        }
    except Exception as e:
        return {"ats_text": f"Error: {str(e)}", "available": True}


def export_cv_to_docx(cv_text: str) -> bytes:
    """Export CV text to a formatted DOCX file. Returns bytes."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)

    # Parse and add content
    lines = cv_text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Detect headings (lines in ALL CAPS or starting with ##)
        if line.startswith("## ") or line.startswith("# "):
            clean = line.lstrip("#").strip().replace("**", "")
            p = doc.add_heading(clean, level=2)
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            clean = line.replace("**", "")
            p = doc.add_heading(clean, level=2)
        elif line == "---":
            # Garis horizontal ASLI (border paragraf), bukan teks "---" mentah
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "CCCCCC")
                pBdr.append(bottom)
                pPr.append(pBdr)
            except Exception:
                pass  # kalau gagal, cukup paragraf kosong - tidak menampilkan "---" mentah
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            p = doc.add_paragraph(style="List Bullet")
            parts = clean.split("**")
            for i, part in enumerate(parts):
                p.add_run(part).bold = (i % 2 == 1)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                p.add_run(part).bold = (i % 2 == 1)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_cv_to_pdf(cv_text: str) -> bytes:
    """Export CV text to a formatted PDF file. Returns bytes."""
    import re
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=40,
        bottomMargin=40,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    heading_style = ParagraphStyle(
        "CVHeading",
        parent=styles["Heading2"],
        fontSize=13,
        spaceAfter=6,
        spaceBefore=14,
        textColor="#1a1a2e",
    )
    body_style = ParagraphStyle(
        "CVBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=4,
        leading=14,
    )
    bullet_style = ParagraphStyle(
        "CVBullet",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=3,
        leftIndent=20,
        leading=14,
        bulletIndent=10,
    )

    elements = []
    lines = cv_text.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 6))
            continue

        # Escape XML special characters for reportlab
        safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if line.startswith("## ") or line.startswith("# "):
            clean = safe_line.lstrip("#").strip()
            clean = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean)
            elements.append(Paragraph(clean, heading_style))
        elif line.upper() == line and len(line) > 3 and not line.startswith("-"):
            clean = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe_line)
            elements.append(Paragraph(clean, heading_style))
        elif line.startswith("- ") or line.startswith("• "):
            clean = safe_line.lstrip("-•").strip()
            clean = re.sub(r'\*\frac{.*?}{.*?}', r'', clean)  # remove math symbols if any
            clean = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean)
            elements.append(Paragraph(f"• {clean}", bullet_style))
        else:
            clean = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe_line)
            elements.append(Paragraph(clean, body_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


MATCH_GAP_PROMPT = """Kamu adalah Career Strategist & ATS Analyst berpengalaman.
Tugasmu adalah membandingkan CV kandidat dengan SATU lowongan kerja spesifik,
lalu berikan analisis kecocokan dan gap yang jujur dan actionable.

Berikan output dalam format berikut (gunakan heading yang sama persis):

## 🎯 Skor Kecocokan: [angka 0-100]/100

## ✅ Skill & Pengalaman yang Match
- [poin yang ada di CV DAN diminta lowongan]
- [poin lain]

## ⚠️ Gap yang Perlu Diperhatikan
- [skill/kualifikasi yang diminta lowongan tapi TIDAK ADA/kurang jelas di CV]
- [poin lain]

## 💡 Strategi Konkret
- [saran spesifik: skill apa yang perlu dipelajari/di-highlight, sertifikasi
  yang relevan, atau cara reframe pengalaman yang ada supaya lebih cocok]
- [saran lain]

## 📊 Ringkasan
[1-2 kalimat kesimpulan jujur: apakah kandidat sudah cukup kompetitif untuk
posisi ini, atau perlu persiapan tambahan]

ATURAN:
1. Skor HARUS jujur dan realistis - jangan selalu tinggi untuk menyenangkan user.
   Kalau memang kurang cocok, skor rendah dan jelaskan kenapa.
2. Gap yang disebut harus SPESIFIK, merujuk ke requirement asli di deskripsi
   lowongan, bukan generik.
3. Jawab dalam Bahasa Indonesia."""


def analyze_match_and_gap(cv_text: str, job_info: dict) -> dict:
    """
    Analyze CV match level against a specific job, and identify skill gaps.

    Args:
        cv_text: User's CV content
        job_info: dict with job_title, company_name, job_description

    Returns dict with:
    - "analysis": AI-generated match/gap analysis markdown
    - "available": whether Gemini is configured
    """
    if not config.is_gemini_configured():
        return {"analysis": None, "available": False}

    try:
        client = config.get_gemini_client()
        target_pos = job_info.get("job_title", "N/A")
        target_company = job_info.get("company_name", "N/A")
        target_desc = job_info.get("job_description", "N/A")

        prompt = (
            f"{MATCH_GAP_PROMPT}\n\n"
            f"[CV KANDIDAT]\n{cv_text}\n\n"
            f"[LOWONGAN SASARAN]\n"
            f"Posisi: {target_pos}\n"
            f"Perusahaan: {target_company}\n"
            f"Deskripsi Lowongan:\n{target_desc[:2000]}"
        )
        response = client.models.generate_content(
            model=config.GEMINI_MODEL,
            contents=prompt,
        )
        return {"analysis": response.text, "available": True}
    except Exception as e:
        return {"analysis": f"Error: {str(e)}", "available": True}
